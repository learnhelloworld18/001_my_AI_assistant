"""Turn a folder of documents into searchable chunks.

    /ingest <path> --collection resume_interview
       |
       v
    discover()   walk, applying the allowlist and the exclusions below
       |
       v
    for each file:  unchanged?  -> skip, costs nothing
       |            changed?    -> drop its old chunks, then add new ones
       v
    load() -> chunk() -> embed -> Chroma, and record the hash

Re-running this is the intended way to keep a collection current. It is not an
"update" command because it does not need to be: unchanged files are free, and
changed ones replace themselves rather than accumulating.

The exclusions are not incidental. Pointed at a real documents folder, a naive
walk pulled in a cloned git repository (201 Python files), stale `_BACKUP_`
copies of live documents, and certificates that are mostly layout - each of
which would compete with genuine prose in retrieval.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from myassistant.rag import manifest, store

log = logging.getLogger("myassistant")

# Only formats we can actually read as text. `.pages` is deliberately absent:
# it is an Apple bundle whose text sits in a proprietary binary format, and its
# zip contains nothing but JPG previews. Export those to PDF instead.
READABLE = {".md", ".txt", ".pdf", ".docx"}

# Directories that are never someone's notes.
SKIP_DIRS = {".git", "node_modules", ".venv", "venv", "__pycache__", ".idea", ".vscode"}

# Filename patterns to skip wherever they appear.
#   *_BACKUP_*  dated copies sitting beside the live document - ingesting both
#               puts a stale and a current answer in retrieval together
#   *Certified* certificates: a name, a date, and a lot of layout
SKIP_PATTERNS = ("*_BACKUP_*", "*Certified*", "*certificate*")

# Exact names to skip - instructions for other tools, not personal content.
SKIP_NAMES = {"CLAUDE.md", "README.md", "ASUtranscript.pdf"}

# Small enough that a retrieved chunk leaves room for the question and the
# answer in a 3B context window; overlapped so a fact split across a boundary
# still appears whole in one of the two chunks.
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# Below this a "successful" parse produced no usable text - a scanned PDF, an
# empty file. Same reasoning as tools/observation.looks_empty: silently storing
# nothing is worse than saying so.
MIN_USEFUL_CHARS = 50


@dataclass
class Result:
    """What one ingest run did. Printed by /ingest, asserted by tests."""

    added: list[str] = field(default_factory=list)  # files (re)ingested
    skipped: list[str] = field(default_factory=list)  # unchanged since last run
    failed: list[tuple[str, str]] = field(default_factory=list)  # (file, why)
    chunks: int = 0

    def summary(self) -> str:
        """One line per outcome, and the failures named rather than swallowed."""
        parts = [f"{len(self.added)} ingested ({self.chunks} chunks)"]
        if self.skipped:
            parts.append(f"{len(self.skipped)} unchanged")
        if self.failed:
            parts.append(f"{len(self.failed)} failed")
        line = ", ".join(parts)
        for name, why in self.failed:
            line += f"\n  ! {Path(name).name}: {why}"
        return line


def _inside_repo(path: Path, root: Path) -> bool:
    """Is this file inside a nested git repository?

    Written as a general rule rather than a hardcoded folder name: a checked-out
    repo under a notes directory is source code that happens to live there, and
    its README describes how to run it, not what you did. Code belongs to
    coding_agent's file tools, which read it precisely and on demand; putting it
    in the vector store only crowds out the prose that answers real questions.
    """
    for parent in [path, *path.parents]:
        if parent == root.parent:
            break
        if (parent / ".git").exists():
            return True
    return False


def _excluded(path: Path, root: Path) -> bool:
    """Every reason to skip a file, in one place."""
    if path.suffix.lower() not in READABLE:
        return True
    if path.name in SKIP_NAMES or path.name.startswith("."):
        return True
    if any(part in SKIP_DIRS for part in path.parts):
        return True
    if any(path.match(pattern) for pattern in SKIP_PATTERNS):
        return True
    return _inside_repo(path, root)


def discover(root: Path) -> list[Path]:
    """Every ingestible file under root, sorted so runs are reproducible."""
    if root.is_file():
        return [] if _excluded(root, root.parent) else [root]
    return sorted(p for p in root.rglob("*") if p.is_file() and not _excluded(p, root))


def _read_pdf(path: Path) -> str:
    """Text layer only. A scanned PDF has none, and comes back empty by design."""
    from pypdf import PdfReader

    return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)


def _read_docx(path: Path) -> str:
    """Paragraphs and table cells - STAR answers are often laid out in tables."""
    import docx

    document = docx.Document(str(path))
    blocks = [p.text for p in document.paragraphs]
    blocks += [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(blocks)


def load(path: Path) -> str:
    """Read one file to plain text. Raises with a readable message on failure."""
    if path.suffix.lower() == ".pdf":
        return _read_pdf(path)
    if path.suffix.lower() == ".docx":
        return _read_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")


def chunk(text: str, source: Path) -> list[Document]:
    """Split into overlapping chunks, each tagged with the file it came from.

    The `source` metadata is load-bearing twice over: it is how drop_source()
    finds a file's old chunks on re-ingest, and it is how an answer can say
    which document it came from.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    return [
        Document(page_content=piece, metadata={"source": str(source), "name": source.name})
        for piece in splitter.split_text(text)
    ]


def ingest(
    root: Path,
    collection: store.Collection,
    *,
    embedding_function: Any | None = None,
    db: Path | None = None,
    force: bool = False,
) -> Result:
    """Ingest a file or a whole tree into one collection.

    `force` re-ingests unchanged files too - useful after changing CHUNK_SIZE,
    where the files are identical but the chunking is not.
    """
    result = Result()
    collection_store = store.get(collection, embedding_function=embedding_function)

    for path in discover(root):
        if not force and manifest.is_unchanged(path, str(collection), db):
            result.skipped.append(str(path))
            continue

        try:
            text = load(path)
        except Exception as e:
            log.exception("ingest failed: %s", path)
            result.failed.append((str(path), f"{type(e).__name__}: {e}"))
            continue

        # An empty parse is a failure worth naming, not a document worth storing.
        # A scanned PDF looks like a success right up until you search for it.
        if len(text.strip()) < MIN_USEFUL_CHARS:
            result.failed.append((str(path), "no extractable text (scanned or empty?)"))
            continue

        pieces = chunk(text, path)

        # Delete before add, always - including on a first ingest, where it is a
        # harmless no-op. Doing it unconditionally means there is no path where
        # a changed file's old chunks survive.
        collection_store.delete(where={"source": str(path)})
        collection_store.add_documents(pieces)
        manifest.record(path, str(collection), db)

        result.added.append(str(path))
        result.chunks += len(pieces)

    return result
